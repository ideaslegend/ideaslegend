from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import markdown2
import requests
import os
from urllib.parse import urlparse
import asyncio
import threading
from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Inches
import cairosvg
import lxml


# 判断一个标签是否为块级元素（可扩展其他块级标签）
def is_block(tag):
    block_tags = ['body','p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
            'ul', 'ol', 'table', 'blockquote', 'pre', 'article', 'aside',
            'canvas', 'details', 'figcaption', 'figure', 'footer', 'header',
            'hr', 'main', 'nav', 'section', 'summary', 'video','form','menu','video','img']
    return tag.name in block_tags

def _set_run_font(run, font_name="新宋体"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def get_current_file_dir():
    """获取当前文件所在目录"""
    return os.path.dirname(os.path.abspath(__file__))

def svg_to_png(svg_path, png_path):
    # 读取并确保编码正确
    with open(svg_path, 'r', encoding='utf-8') as f:
        svg_content = f.read()

    soup = BeautifulSoup(svg_content, 'xml')
    body = soup.body if soup.body else soup
    for g_element in body.select('g'):  # 选择所有<g>元素
        if 'font-family' in g_element.attrs:
            print(g_element['font-family'])
            g_element['font-family'] = 'Microsoft YaHei,MS Mincho,Malgun Gothic,Nirmala UI,Verdana,Geneva,DejaVu Sans,sans-serif'  
            print(g_element['font-family'])
    content = body.decode_contents()
    try:
        cairosvg.svg2png(bytestring=content.encode('utf-8'), write_to=png_path,dpi=300)
        #cairosvg.svg2png(url=svg_path, write_to=png_path,dpi=300)
    except Exception as e:
        print(f"svg图片处理失败: {str(e)}")
        
def download_image(image_url, save_folder='images'):
    """
    根据图片链接下载图片，并返回下载后的图片本地地址。

    :param image_url: 图片的URL链接
    :param save_folder: 保存图片的文件夹，默认为'images'
    :return: 下载后的图片本地地址
    """
    try:
        
        response = requests.get(image_url, stream=True) 
        response.raise_for_status()

        # 解析URL以获取文件名和扩展名
        parsed_url = urlparse(image_url)
        file_name = os.path.basename(parsed_url.path)
        
        # 尝试从Content-Type获取图片格式
        content_type = response.headers.get('Content-Type', '')
        if 'image/' in content_type or 'application/svg+xml' in content_type:
            ext = content_type.split('/')[-1].lower()
            if ext == 'svg+xml':
                ext = 'svg'
        else:
            # 如果Content-Type不可用，尝试从URL中获取扩展名
            ext = file_name.rsplit('.', 1)[1].lower() if '.' in file_name else 'jpg'
        
        # 确保扩展名是有效的图片格式
        valid_extensions = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp', 'svg'}
        if ext not in valid_extensions:
            # 如果是SVG格式的内容，强制使用svg扩展名
            if b'<svg' in response.content[:1024]:
                ext = 'svg'
            else:
                ext = 'jpg'  # 默认使用jpg格式
        
        # 处理文件名
        if '.' in file_name:
            file_name = file_name.rsplit('.', 1)[0]
        file_name = f"{file_name}.{ext}"

        # 创建保存文件夹（如果不存在）
        if not os.path.exists(save_folder):
            os.makedirs(save_folder)

        # 拼接保存路径
        save_path = os.path.join(save_folder, file_name)
        # 将图片数据写入文件
        with open(save_path, 'wb') as file:
            for chunk in response.iter_content(chunk_size=8192):
                file.write(chunk)

        return save_path

    except Exception as e:
        print(f"下载图片时发生错误: {e}")
        return None
def doc_image_insert(doc, md_path, elem):
    src = elem.get('src')
    if src:
        if src.startswith('http://') or src.startswith('https://'):
            # 使用改进后的download_image函数下载图片
            img_path = download_image(src, get_current_file_dir() + '/temp/images')
            if not img_path:
                p = doc.add_paragraph(f"[图片下载失败: {src}]")
                for run in p.runs:
                    _set_run_font(run)
                return
        else:
            img_path = (md_path.parent/src).resolve()
        
        if os.path.exists(img_path):
            try:
                # 获取文档可用宽度
                section = doc.sections[0]
                available_width = section.page_width - section.left_margin - section.right_margin
                
                # 如果是SVG格式，需要先转换为PNG
                if img_path.endswith('.svg'):
                    png_path = img_path.rsplit('.', 1)[0] + '.png'
                    svg_to_png(img_path, png_path)
                    img_path = png_path
                
                
                doc.add_picture(str(img_path), width=available_width)
                
                # 如果是临时转换的PNG，删除它
                if img_path.endswith('.png') and src.endswith('.svg'):
                    os.remove(img_path)
                    
            except Exception as e:
                p = doc.add_paragraph(f"[图片处理失败: {str(e)}]")
                print(f"图片处理失败: {str(e)}")
                
        else:
            p = doc.add_paragraph(f"[图片文件不存在: {img_path}]")
            for run in p.runs:
                _set_run_font(run)
        
def _doc_image_process(doc,md_path,elem):    
    """
    图片处理函数
    :param doc: doc元素
    :param md_path: md文件路径
    :param elem: 图片元素
    """
    src = elem.get('src')
    if src:
        if src.startswith('http://') or src.startswith('https://'):
            import queue
            def thread_function(queue):
                path = download_image(src,get_current_file_dir()+ '/temp'+'/images')
                queue.put(path)  
            q = queue.Queue()
            thread = threading.Thread(target=thread_function,args=(q,))
            thread.start()
            thread.join(5)
            img_path = q.get() if not q.empty() else None
            
            #img_path = download_image(src,get_current_file_dir()+ '/temp'+'/images')
            if not img_path:
                p = doc.add_paragraph(f"[图片下载失败: {src}]")
                for run in p.runs:
                    _set_run_font(run)
        else:
            img_path = (md_path.parent/src).resolve()
        
        try:
            section = doc.sections[0]
            available_width = section.page_width - section.left_margin - section.right_margin
            if str(img_path).endswith('.svg'):
                    png_path = img_path.rsplit('.', 1)[0] + '.png'
                    svg_to_png(img_path, png_path)
                    img_path = png_path
            doc.add_picture(str(img_path), width=available_width)
        except Exception as e:
            p = doc.add_paragraph(f"[图片加载失败: {src}]")
            for run in p.runs:
                _set_run_font(run)
def get_img_src(element):
    """
    检查元素内部是否包含img标签并返回其src属性
    参数:
    element: 要检查的HTML元素对象
    返回:
    str: 如果找到img标签则返回其src属性值，否则返回None
    """
    # 检查元素本身是否是img标签
    if hasattr(element, 'name') and element.name == 'img':
        return element.get('src', None)
        
    # 检查元素内部是否包含img标签
    img_tags = element.find_all('img')
    if img_tags:
        return img_tags[0].get('src', None)
        
    return None
# 添加超链接——使用 python-docx 的底层 API构造一个超链接对象
def add_hyperlink(paragraph,element, url, text):
    part = paragraph.part
    
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    

    # 创建 run 元素及其属性（下面设置了超链接默认风格：蓝色+下划线）
    img_url = get_img_src(element)
    hyperlink = OxmlElement('w:hyperlink')
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_rStyle = OxmlElement('w:rStyle')
    new_rStyle.set(qn('w:val'), 'Hyperlink')
    # 设置字体颜色为蓝色（RGB 0000FF）
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    rPr.append(new_rStyle)
    new_run.append(rPr)
    if  img_url is not None:
        try:
            if img_url.startswith('http://') or img_url.startswith('https://'):
                img_dir = os.path.join(get_current_file_dir(), 'temp', 'images')
                os.makedirs(img_dir, exist_ok=True)
                
                img_path = download_image(img_url, img_dir)
            else:
                strPsth = str(get_current_file_dir()) 
                my_path = Path(strPsth)
                img_path = (my_path/img_url).resolve()
            # 创建图片运行对象
            run = paragraph.add_run()
            if str(img_path).endswith('.svg'):
                    png_path = img_path.rsplit('.', 1)[0] + '.png'
                    svg_to_png(img_path, png_path)
                    img_path = png_path
            # 添加图片
            run.add_picture(str(img_path),width=Inches(4))
            
            # 创建超链接
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            # 将超链接添加到图片
            #hyperlink.append(run._element)
            new_run.append(run._element)
            
        except UnrecognizedImageError:
            # 处理不支持的图片格式
            paragraph.add_run(f"[不支持的图片格式: {img_path}]")    
        except Exception as e:
            paragraph.add_run(f"[图片处理失败: {str(e)}]")
    else:
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
   
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# 处理表格：遍历所有行和单元格
def process_table(element, doc):
    # 创建Word表格
    row_count = len(element.find_all('tr'))
    col_count = max(len(tr.find_all(['td','th'])) for tr in element.find_all('tr'))
    table = doc.add_table(rows=row_count, cols=col_count)
    
    # 设置表格样式
    table.style = 'Table Grid'
    #table.style.font.size = Pt(12)
    table.autofit = True

    # 填充表格数据
    for row_idx, tr in enumerate(element.find_all('tr')):
        cells = tr.find_all(['td', 'th'])
        for col_idx, td in enumerate(cells):
            cell = table.cell(row_idx, col_idx)  # 先获取单元格
            
            # 创建段落并应用字体
            p = cell.paragraphs[0]
            run = p.add_run(td.get_text().strip())
            
            # 首行样式处理
            if row_idx == 0:
                run.font.bold = True
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'ADD8E6')
                cell._element.get_or_add_tcPr().append(shading)
            
            # 继承段落对齐方式
            #'text-align:center;'
            if 'style' in td.attrs:
                align_map = {
                    'text-align:left;': WD_ALIGN_PARAGRAPH.LEFT,
                    'text-align:center;': WD_ALIGN_PARAGRAPH.CENTER,
                    'text-align:right;': WD_ALIGN_PARAGRAPH.RIGHT
                }
                cell.paragraphs[0].alignment = align_map.get(td['style'].lower(), WD_ALIGN_PARAGRAPH.LEFT)
                # 处理单元格内超链接
            for link in td.find_all('a'):
                add_hyperlink(cell.paragraphs[0],td, link.get('href', ''), link.get_text())

# 直接处理块级元素内联子元素，parameters 中可累积格式（如 bold、italic、underline）
def direct_process_inline_children(md_file_path,document,soup_element, paragraph, formatting=None):
        if formatting is None:
            formatting = {}
        if isinstance(soup_element, NavigableString):
            text = str(soup_element)
            if text.strip():
                run = paragraph.add_run(text)
            if formatting.get('bold'):
                run.bold = True
            if formatting.get('italic'):
                run.italic = True
            if formatting.get('underline'):
                run.underline = True
        elif hasattr(soup_element, 'name'):
            print(soup_element.name)
            # 根据标签更新格式信息
            if soup_element.name in ['strong', 'b']:
                new_formatting = formatting.copy()
                new_formatting['bold'] = True
                process_inline_children(md_file_path,document,soup_element, paragraph, new_formatting)
            elif soup_element.name in ['em', 'i']:
                new_formatting = formatting.copy()
                new_formatting['italic'] = True
                process_inline_children(md_file_path,document,soup_element, paragraph, new_formatting)
            elif soup_element.name == 'u':
                new_formatting = formatting.copy()
                new_formatting['underline'] = True
                process_inline_children(md_file_path,document,soup_element, paragraph, new_formatting)
            elif soup_element.name == 'a':
                # 对超链接标签，提取 href 和文本并调用 add_hyperlink
                url = soup_element.get('href', '')
                hyperlink_text = "".join(soup_element.stripped_strings)
                add_hyperlink(paragraph,soup_element, url, hyperlink_text)
            elif soup_element.name == 'img':
                _doc_image_process(document,md_file_path,soup_element)
                #doc_image_insert(document,md_file_path,soup_element)
            elif soup_element.name == 'code':
                # 处理代码标签
                new_formatting = formatting.copy()
                new_formatting['underline'] = True
                process_inline_children(md_file_path,document,soup_element, paragraph, new_formatting)
            elif soup_element.name == 'table':
                p = document.add_paragraph().add_run()
                process_table(soup_element, p)
            else:
                # 对于其它标签，如果不是块级，则递归调用内联处理；若为块级则单独作为一个块处理
                if not is_block(soup_element):
                    process_inline_children(md_file_path,document,soup_element, paragraph, formatting)
                else:
                    process_element(md_file_path,soup_element,document)
# 递归处理内联子元素，parameters 中可累积格式（如 bold、italic、underline）
def process_inline_children(md_file_path,document,soup_element, paragraph, formatting=None):
    if formatting is None:
        formatting = {}
    print(soup_element.name)
    
    for child in soup_element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip():
                run = paragraph.add_run(text)
            if formatting.get('bold'):
                run.bold = True
            if formatting.get('italic'):
                run.italic = True
            if formatting.get('underline'):
                run.underline = True
        elif hasattr(child, 'name'):
            print(child.name)
            # 根据标签更新格式信息
            if child.name in ['strong', 'b']:
                new_formatting = formatting.copy()
                new_formatting['bold'] = True
                process_inline_children(md_file_path,document,child, paragraph, new_formatting)
            elif child.name in ['em', 'i']:
                new_formatting = formatting.copy()
                new_formatting['italic'] = True
                process_inline_children(md_file_path,document,child, paragraph, new_formatting)
            elif child.name == 'u':
                new_formatting = formatting.copy()
                new_formatting['underline'] = True
                process_inline_children(md_file_path,document,child, paragraph, new_formatting)
            elif child.name == 'a':
                # 对超链接标签，提取 href 和文本并调用 add_hyperlink
                url = child.get('href', '')
                hyperlink_text = "".join(child.stripped_strings)
                add_hyperlink(paragraph,child, url, hyperlink_text)
            elif child.name == 'img':
                _doc_image_process(document,md_file_path,child)
                #doc_image_insert(document,md_file_path,soup_element)
            elif child.name == 'code':
                # 处理代码标签
                new_formatting = formatting.copy()
                new_formatting['underline'] = True
                process_inline_children(md_file_path,document,child, paragraph, new_formatting)
            elif child.name == 'table':
                p = document.add_paragraph()
                process_table(child, document)
            else:
                # 对于其它标签，如果不是块级，则递归调用内联处理；若为块级则单独作为一个块处理
                if not is_block(child):
                    process_inline_children(md_file_path,document,child, paragraph, formatting)
                else:
                    process_element(md_file_path,child,document)
        else:
            # 对于其它类型元素，按内联处理
            direct_process_inline_children(md_file_path,document,child, paragraph, formatting)
    
# 处理块级标签（标题、段落、div、列表、表格等）。
def process_element(md_file_path,element,doc,main_p=None):
    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(element.name[1])
        p = doc.add_heading("", level=level)
        process_inline_children(md_file_path,doc,element, p)
        if element.get('align') == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    elif element.name == 'p':
        p = doc.add_paragraph()
        for child in element.children:
            if isinstance(child, NavigableString):
                text = child.strip()
                p.add_run(text)
            elif hasattr(child, 'name') and is_block(child):
                process_element(md_file_path,child, doc)
            elif hasattr(child, 'name')and not is_block(child):
                print(child.name)
                direct_process_inline_children(md_file_path,doc,child, p)
            else:
                # 当作内联元素处理
                process_inline_children(md_file_path,doc,child, p)
        if element.get('align') == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif element.name == 'div':
        # 对 div 逐个处理子节点。如果子节点为块级则单独处理；内联节点聚合到一起。
        p = doc.add_paragraph()
        
        for child in element.children:
            if isinstance(child, NavigableString):
                text = child.strip()
                if text:
                    p.add_run(text)
            elif hasattr(child, 'name') and is_block(child):
                process_element(md_file_path,child, doc)
            elif hasattr(child, 'name')and not is_block(child):
                print(child.name)
                direct_process_inline_children(md_file_path,doc,child, p)
            else:
                # 当作内联元素处理
                process_inline_children(md_file_path,doc,child, p)
        if element.get('align') == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif element.name == 'ul':
        # 处理无序列表
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph("", style='List Bullet')
            process_inline_children(md_file_path,doc,li, p)
    elif element.name == 'ol':
        # 处理有序列表
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph("", style='List Number')
            process_inline_children(md_file_path,doc,li, p)
    elif element.name == 'table':
        p = doc.add_paragraph()
        process_table(element, doc)
    elif element.name == 'blockquote':
        p = doc.add_paragraph()
        for child in element.children:
            if isinstance(child, NavigableString):
                text = child.strip()
                if text:
                    p.add_run(text)
            elif hasattr(child, 'name') and is_block(child):
                process_element(md_file_path,child, doc)
            elif hasattr(child, 'name')and not is_block(child):
                
                direct_process_inline_children(md_file_path,doc,child, p)
            else:
                # 当作内联元素处理
                process_inline_children(md_file_path,doc,child, p)
        if element.get('align') == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif element.name == 'img':
        #doc_image_insert(doc,md_file_path,element)
        _doc_image_process(doc,md_file_path,element)
    else:
        # 其它块级或不常见标签按内联来处理
        if main_p is None:
            main_p = doc.add_paragraph()
        direct_process_inline_children(md_file_path,doc,element,main_p)


async def test_md_to_docx(md_path: Path, output_path: Path):
    md_text = md_path.read_text(encoding='utf-8')
    md_text = '\n'.join(line for line in md_text.splitlines() if line.strip())  # 删除空行
    html = markdown2.markdown(md_text, extras=[ "fenced-code","toc","tables"])  #
   
    soup = BeautifulSoup(html, 'lxml')    
    doc = Document()
    body = soup.body if soup.body else soup
    
   
    # for img_li in body.find_all('img'):
    #     #  path = download_image(img_li.get('src'), get_current_file_dir() + '/temp/images')
    #     #  doc.add_picture(str(path), width=Inches(4))
    #     doc_image_insert(doc,md_path,img_li)
    
    main_p = doc.add_paragraph()
    for child in body.children:
        if hasattr(child, 'name') and is_block(child):
            print(child.name)
            process_element(md_path,child,doc,main_p)
        elif hasattr(child, 'name')and not is_block(child):
            print(child.name)
            direct_process_inline_children(md_path,doc,child,main_p)
        elif isinstance(child, NavigableString):
            text = child.strip()
            if text:
                main_p.add_run(text)
                #doc.add_paragraph(text)

    # 设置文档样式
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.save(str(output_path))
 
async def md_to_docx(path_or_content, output_path: Path):

    if isinstance(path_or_content, Path):
        md_text = path_or_content.read_text(encoding='utf-8')
        md_path = path_or_content
    elif isinstance(path_or_content, str):
        md_text = path_or_content
        md_path = Path()
    md_text = '\n'.join(line for line in md_text.splitlines() if line.strip())  # 删除空行
    html = markdown2.markdown(md_text, extras=[ "fenced-code","toc","tables"])  #
   
    soup = BeautifulSoup(html, 'lxml')    
    doc = Document()
    body = soup.body if soup.body else soup
    # main_p = doc.add_paragraph()
    # for child in body.children:
    #     if hasattr(child, 'name') and is_block(child):
    #         print(child.name)
    #         if child.name == 'strong':
    #             print("strong")
    #         process_element(md_path,child,doc,main_p)
    #     elif hasattr(child, 'name')and not is_block(child):
    #         print(child.name)
    #         if child.name == 'strong':
    #             print("strong")
    #         direct_process_inline_children(md_path,doc,child,main_p)
    #     elif isinstance(child, NavigableString):
    #         text = child.strip()
    #         if text:
    #             main_p.add_run(text)
    #             #doc.add_paragraph(text)
    
    for child in body.children:
        if hasattr(child, 'name'):
            print(child.name)
            if child.name == 'strong':
                print("strong")
            process_element(md_path,child,doc)
        elif isinstance(child, NavigableString):
            text = child.strip()
            if text:
                doc.add_paragraph(text)

    # 设置文档样式
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.save(str(output_path))

def main():
    md_path = Path("./common/README_2.md")
    output_path = Path('./common/test1.docx')
    # md_path = Path("README_2.md")
    # output_path = Path('test1.docx')
    #asyncio.run(test_md_to_docx(md_path, output_path))


    content = """"""
    asyncio.run(md_to_docx(md_path, output_path))
    pass

if __name__ == '__main__':
    main()
    #svg_to_png('./common/temp/images/বাংলা-d9d9d9.svg', './common/temp/images/test.png')